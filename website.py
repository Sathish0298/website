from dotenv import load_dotenv
from langchain.prompts import PromptTemplate
import docx
import pandas as pd
from langchain_openai import ChatOpenAI
from sklearn.feature_extraction.text import CountVectorizer
import matplotlib.pyplot as plt
import seaborn as sns
from wordcloud import WordCloud
import streamlit as st
from langchain_community.document_loaders import TextLoader
from io import StringIO
import requests
from io import BytesIO
from bs4 import BeautifulSoup
from urllib.parse import urlparse
import re
import time


def load_data():
    data = pd.read_csv("Data/cleaned_reviews.csv")
    data["feedback"] = data["Rating"].apply(lambda x: 1 if x > 2 else 0)
    data.dropna(inplace=True)
    data["length"] = data["Comments"].apply(len)
    cv = CountVectorizer(stop_words="english")

    # Combine all reviews for each feedback category and splitting them into individual words
    neg_reviews = " ".join(
        [review for review in data[data["feedback"] == 0]["Comments"]]
    )
    neg_reviews = neg_reviews.lower().split()

    pos_reviews = " ".join(
        [review for review in data[data["feedback"] == 1]["Comments"]]
    )
    pos_reviews = pos_reviews.lower().split()

    # Finding words from reviews which are present in that feedback category only
    unique_negative = [x for x in neg_reviews if x not in pos_reviews]
    unique_negative = " ".join(unique_negative)

    unique_positive = [x for x in pos_reviews if x not in neg_reviews]
    unique_positive = " ".join(unique_positive)

    # Assume unique_negative is a string containing all unique negative words separated by spaces
    with open("unique_negative_keywords.txt", "w") as file:
        file.write(unique_negative)
    return data


# Visualization functions
def plot_rating_distribution(data):
    fig, ax = plt.subplots()
    data["Rating"].value_counts().plot.bar(color="red", ax=ax)
    plt.title("Rating Distribution Count")
    plt.xlabel("ratings")
    plt.ylabel("Count")
    st.markdown("#### Rating Distribution graph", unsafe_allow_html=True)
    buf = BytesIO()
    fig.savefig(buf, format="png")
    st.image(buf)


def plot_wordcloud(data):
    text = " ".join(review for review in data["Comments"])
    wordcloud = WordCloud(background_color="white").generate(text)
    fig = plt.figure(figsize=(10, 5))
    plt.imshow(wordcloud, interpolation="bilinear")
    plt.axis("off")
    st.markdown("#### Key Phrases Word Cloud", unsafe_allow_html=True)
    buf = BytesIO()
    fig.savefig(buf, format="png")
    st.image(buf)


def generate_improvement_suggestions(texts):

    template_string = """ You are an AI language model trained to analyze unique key phrases or words from the negative reviews of a product and generate snew product design for product improvements. \
    The unique key phrases or words from the negative reviews of a product are: {document} \

    Based on the following Unique key phrases or words from the negative reviews, suggest top new product design improvements:
    Make sure the improvements are generic and more specific to a particular issue or key phrase.
    Don't suggest similar improvements more than once.
    Each Improvements suggested should have following subdivisions:
    Improvement
    corresponing key phrase or word
    Impact
    Steps to Implement
    (steps to implement should be in bullet points)
    """

    prompt = PromptTemplate(
        template=template_string,
        input_variables=["document"],
    )

    model = ChatOpenAI(
        model="gpt-3.5-turbo",
        max_tokens=3000,
        n=1,
        temperature=0.7,
    )

    chain = prompt | model

    completion = chain.invoke({"document": texts})

    return completion.content


def read_word_document(file_path):
    doc = docx.Document(file_path)
    improvements = []
    temp_dict = {}
    for para in doc.paragraphs:
        text = para.text.strip()
        if text.startswith("Improvement"):
            if temp_dict:
                improvements.append(temp_dict)
            temp_dict = {"Improvement": text}
        elif text.startswith("Corresponding key phrase or word"):
            temp_dict["Key Phrase"] = text
        elif text.startswith("Impact"):
            temp_dict["Impact"] = text
        elif text.startswith("Steps to Implement"):
            temp_dict["Steps"] = []
        elif text.startswith("-"):
            temp_dict["Steps"].append(text)
    if temp_dict:
        improvements.append(temp_dict)
    return improvements


def display_improvements():
    doc_path = "outputs/full_improvement_suggestions.docx"
    try:
        with open(doc_path, "rb") as file:
            btn = st.download_button(
                label="Download Improvement Suggestions",
                data=file,
                file_name="improvement_suggestions.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
    except Exception as e:
        st.error(f"Failed to prepare the document for download: {str(e)}")


# Function to determine the website based on the URL
def determine_website(url):
    parsed_url = urlparse(url)
    hostname = parsed_url.hostname

    if "amazon" in hostname:
        return "amazon"
    elif "flipkart" in hostname:
        print("Flipkart")
        return "flipkart"
    elif "snapdeal" in hostname:
        return "snapdeal"
    else:
        return None


# Amazon review scraping function
def amazon_review_scraper(url, page):
    reviews = []
    with requests.Session() as session:
        # Set a user agent to mimic a web browser
        session.headers.update(
            {
                "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/121.0.0.0 Safari/537.36 Edg/121.0.0.0"
            }
        )
        url = f"{url}&pageNumber={page}"
        print(f"Accessing {url}")  # Debugging statement
        response = session.get(url)
        if response.url != url:
            print(f"Redirected to {response.url} instead of {url}")
        soup = BeautifulSoup(response.content, "lxml")

        for review in soup.find_all("div", {"class": "a-section review aok-relative"}):
            name_element = review.find("span", {"class": "a-profile-name"})
            rating_element = review.find("i", {"data-hook": "review-star-rating"})
            comments_element = review.find(
                "div", {"class": "a-row a-spacing-small review-data"}
            )

            if name_element and rating_element and comments_element:
                rating_text = rating_element.text.strip()
                rating = int(float(rating_text.split(" ")[0]))
                review_data = {
                    "Name": name_element.text.strip(),
                    "Rating": rating,
                    "Comments": comments_element.text.strip(),
                }
            reviews.append(review_data)

            time.sleep(1)

    return reviews


# Flipkart review scraping function
def flipkart_review_scraper(url, page):
    url = f"{url}{page}"

    reviews = []

    user_agent = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/121.0.0.0 Safari/537.36 Edg/121.0.0.0"
    }

    response = requests.get(url=url, headers=user_agent)
    soup = BeautifulSoup(response.content, "html.parser")

    for review in soup.find_all(
        "div", {"class": "_27M-vq"}
    ):  # Updated class for review container
        name = review.find(
            "p", {"class": "_2sc7ZR _2V5EHH"}
        ).text  # Updated class for reviewer's name
        rating = review.find(
            "div", {"class": "_3LWZlK"}
        ).text  # Updated class for rating
        comments = review.find(
            "div", {"class": "t-ZTKy"}
        ).text  # Updated class for review comments
        comments = re.sub(
            r"\s*READ\s+MORE\s*", "", comments
        )  # Remove 'READ MORE' links if present

        data = {
            "Name": name,
            "Rating": rating,
            "Comments": comments.strip(),
        }

        reviews.append(data)

    return reviews


# Snapdeal review scraping function
def snapdeal_review_scraper(url, page):
    url = f"{url}{page}"
    reviews = []

    user_agent = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/121.0.0.0 Safari/537.36 Edg/121.0.0.0"
    }

    response = requests.get(url=url, headers=user_agent)
    soup = BeautifulSoup(response.content, "html.parser")

    reviews_skipped = 0  # Counter to track the number of reviews skipped

    for review in soup.find_all("div", {"class": "user-review"}):
        if reviews_skipped < 2:
            reviews_skipped += 1
            continue  # Skip the first two reviews

        rating_elements = review.find_all("i", class_="sd-icon sd-icon-star active")
        rating = len(rating_elements)
        name = review.find("div", {"class": "_reviewUserName"}).get("title")
        comments = review.find("p").text

        data = {
            "Name": name,
            "Rating": rating,
            "Comments": comments,
        }

        reviews.append(data)

    return reviews


def main():
    st.set_page_config(page_title="DaD BOT", layout="wide")

    # Custom CSS to style the Streamlit app
    st.markdown(
        """
    <style>
    .big-font {
        font-size:30px !important;
        font-weight: bold;
        color: #0c6efc;
    }
    .button-style {
        display: inline-flex;
        align-items: center;
        justify-content: center;
        width: 100%;
        height: 3rem;
        padding: .5rem;
        margin: .5rem 0;
        background-color: #4CAF50; /* Green background */
        color: white; /* White text */
        border-radius: 5px;
        border: none;
        font-size: 18px;
        cursor: pointer;
        transition: background-color 0.3s;
    }
    .button-style:hover {
        background-color: #45a049; /* Darker green background on hover */
    }
    .streamlit-input {
        margin-bottom: 20px;
    }
    .streamlit-header {
        margin-bottom: 5px;
        color: #0c6efc;
    }
    </style>
""",
        unsafe_allow_html=True,
    )

    # Header
    st.markdown('<p class="big-font">DaD BOT</p>', unsafe_allow_html=True)

    with st.form("product_link_form"):
        product_link = st.text_input("Paste the link to the product:")
        submit_link = st.form_submit_button("Get the Product Review")
        if submit_link and product_link:
            initiate_scraping(product_link)

    # Display and Download Scraped Data Section
    st.header("View and Download Scraped Data")
    if st.button("View Scraped Data"):
        display_csv()
    download_csv()

    # Product Specification Section
    with st.container():
        st.markdown("## Display the new product specification", unsafe_allow_html=True)
        if st.button("View The Specification", key="product_spec"):
            view_product_specification()


def scrape_reviews(urls):
    url_suffix = ""
    scraper_func = None
    all_reviews = []
    for url in urls:
        website = determine_website(url)
        if not website:
            print(f"Unsupported website for URL: {url}")
            continue

        print(f"Scraping {website} reviews from {url}")

        if website == "amazon":
            scraper_func = amazon_review_scraper
            url_suffix = "?th=1&pageNumber="
        elif website == "flipkart":
            scraper_func = flipkart_review_scraper
            url_suffix = "&page="
        elif website == "snapdeal":
            scraper_func = snapdeal_review_scraper
            url_suffix = "?page="

        # Assuming we are scraping only the first 3 pages for demonstration
        for page in range(1, 3):
            full_url = f"{url}{url_suffix}{page}"
            reviews = scraper_func(full_url, page)
            all_reviews.extend(reviews)

    # Convert list of reviews to a DataFrame and save to CSV string
    reviews_df = pd.DataFrame(all_reviews)
    reviews_df.to_csv("Data/cleaned_reviews.csv", index=False)
    csv_string = reviews_df.to_csv(index=False)
    return csv_string


# Function to initiate scraping and save the results
def initiate_scraping(url):
    reviews_csv = scrape_reviews([url])
    st.session_state["reviews_csv"] = reviews_csv

    st.success("Scraped data has been updated!")


# Function to display the CSV content
def display_csv():
    data = load_data()

    st.title("Data Visualization")

    plot_rating_distribution(data)

    plot_wordcloud(data)

    if "reviews_csv" in st.session_state and st.session_state["reviews_csv"]:
        df = pd.read_csv(StringIO(st.session_state["reviews_csv"]))
        st.dataframe(df.head(10))
    else:
        st.error("No data available. Please scrape some data first.")


# Function to download the CSV file
def download_csv():
    if "reviews_csv" in st.session_state and st.session_state["reviews_csv"]:
        st.download_button(
            label="Download CSV",
            data=st.session_state["reviews_csv"],
            file_name="Cleaned_reviews.csv",
            mime="text/csv",
        )
    else:
        st.error("No data available to download.")


def view_product_specification():

    load_dotenv()

    loader = TextLoader("unique_negative_keywords.txt")
    document = loader.load()

    improvement_suggestions = generate_improvement_suggestions(document)

    print(improvement_suggestions)

    df_improvement_suggestions = pd.DataFrame()
    list_improvement_suggestions = []
    list_improvement_suggestions.append(improvement_suggestions)
    df_improvement_suggestions["pros_cons"] = list_improvement_suggestions
    output_file_improvement_suggestions = "outputs/full_improvement_suggestions.xlsx"
    df_improvement_suggestions.to_excel(
        output_file_improvement_suggestions, index=False
    )

    output_file_improvement_suggestions = "outputs/full_improvement_suggestions.docx"
    doc = docx.Document()

    table = doc.add_table(rows=1, cols=1)
    header_cells = table.rows[0].cells
    header_cells[0].text = "Improvement Suggestions"

    for index, row in df_improvement_suggestions.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row["pros_cons"])

    doc.save(output_file_improvement_suggestions)
    st.title("Product Improvement Suggestions")
    display_improvements()

    doc_path = "outputs/full_improvement_suggestions.docx"

    try:
        doc = docx.Document(doc_path)
        if doc.tables:
            for table_index, table in enumerate(doc.tables):
                # Create a DataFrame to store the table data
                data = []
                for row in table.rows:
                    # Preserve extra spaces by using 'pre' tag in HTML for each cell
                    row_data = [
                        "<pre>" + cell.text.replace("\n", "<br>") + "</pre>"
                        for cell in row.cells
                    ]
                    data.append(row_data)

                # Convert list to DataFrame
                df = pd.DataFrame(data)

                if len(df) > 1:  # Check to ensure there's more than one row
                    df.columns = df.iloc[0]  # Set first row as header
                    df = df[1:]  # Remove the first row after setting it as header

                # Display the DataFrame in Streamlit using HTML to preserve spacing
                st.markdown(f"## Table {table_index + 1}")
                st.write(df.to_html(escape=False, index=False), unsafe_allow_html=True)
        else:
            st.error("No tables found in the document.")
    except Exception as e:
        st.error(f"Failed to read the document: {str(e)}")


if __name__ == "__main__":
    main()
